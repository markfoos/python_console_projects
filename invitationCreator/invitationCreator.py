#invitation creator, makes a word document that has an invitiation
#Created by Mark Foos

import docx

f = open('names.txt', 'r+')

namesRaw = [line for line in f.readlines()]
namesClean = [x.replace('\n','') for x in namesRaw]
f.close()


print("Invitiations will be created for: ")
print(namesClean)


print("Creating Word Document")
doc = docx.Document()
for i in namesClean:
    print("creating invitation for " + i)

    doc.add_heading("Cave Story Party!", 0)
    doc.add_picture('Cave_Story_title_screen.png', width=docx.shared.Inches(4), height
=docx.shared.Inches(3))
    doc.add_paragraph(i + ", You are cordially invited to attend the first ever annual Cave Story Party")
    doc.add_paragraph("It will take place at Mr. Barnes House in Harrison Arkansas on the date of the Apocolypse")
    doc.add_paragraph("See you there!")

    doc.add_page_break()

doc.save('InvitationCreator.docx')
print("Saved word Document")
